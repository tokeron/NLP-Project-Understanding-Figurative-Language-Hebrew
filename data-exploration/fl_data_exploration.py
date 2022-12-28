from datasets import load_dataset, load_metric
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.enum.text import WD_COLOR_INDEX
from box import Box
import yaml
from wordcloud import WordCloud, STOPWORDS
from bidi.algorithm import get_display
import os
from highlight_text import HighlightText, ax_text, fig_text

os.environ['KMP_DUPLICATE_LIB_OK']='True'

label_names = {"O": 0, "B-metaphor": 1, "I-metaphor": 2}


def print_predictions(words_list, labels_list, predictions_list, filename):
    """
    :param words_list: List of words
    :param labels_list: List of labels
    :param predictions_list: List of predictions
    :param filename: Filename to save the output
    """
    doc = docx.Document()
    doc.add_heading('Words with metaphors highlighted', 0)
    for index, (words, labels, predictions) in enumerate(zip(words_list, labels_list, predictions_list)):
        para = doc.add_paragraph('Example: ' + str(index) + '\n')
        for word, label, prediction in zip(words, labels, predictions):
            if label == 0 and prediction == 0:
                para.add_run(word + ' ')
            elif label != 0 and prediction != 0:  # green for correct prediction
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GREEN
            elif label != 0 and prediction == 0:  # red for metaphor that was not predicted
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.RED
            elif label == 0 and prediction != 0:  # pink for not metaphor that was predicted (wrongly)
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.PINK
    # save the document
    doc.save(filename)


def print_labels(words_list, labels_list, filename):
    """
    :param words_list: List of words
    :param labels_list: List of labels
    :param filename: Output filename
    """
    doc = docx.Document()
    doc.add_heading('Words with metaphors highlighted', 0)
    for index, (words, labels) in enumerate(zip(words_list, labels_list)):
        para = doc.add_paragraph('Example: ' + str(index) + '\n')
        for word, label in zip(words, labels):
            if label == 0:
                para.add_run(word + ' ')
            elif label == 1:
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.BRIGHT_GREEN
            else:  # label == 2
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GREEN
    # Now save the document
    doc.save(filename)


def plot_statistics(label_hist, word_dict, examples_with_metaphor):
    """
    :param label_hist: Label histogram
    :param word_dict: Word dictionary
    :param examples_with_metaphor: Number of examples with metaphor
    """
    plt.bar(range(3), label_hist)
    # show numbers on the bars
    for i in range(3):
        plt.text(i, label_hist[i], str(label_hist[i]))
    # change sticks names
    plt.xticks(range(3), [label for label in label_names])
    plt.xlabel('Label')
    plt.ylabel('Count')
    plt.title('Label Distribution')
    plt.show()

    # Number of examples with at least one metaphorx
    plt.figure(figsize=(10, 5))
    plt.bar(range(2), examples_with_metaphor, align='center')
    # show numbers on the bars
    for i in range(2):
        plt.text(i, examples_with_metaphor[i], str(examples_with_metaphor[i]))
    plt.xticks(range(2), ['Without metaphor', 'With metaphor'])
    plt.xlabel('Label')
    plt.ylabel('Number of examples')
    plt.title('Number of examples with at least one label')
    plt.show()

    # create a dataframe from the dictionary
    df = pd.DataFrame.from_dict(word_dict, orient='index')
    df.rename(columns={0: "O", 1: "B", 2: "I"}, inplace=True)
    df.index.name = 'word'

    # Plot top 10 words with highest frequency of O label
    plt.figure(figsize=(10, 10))
    df.sort_values('O', ascending=False).head(10).plot.bar(y=['O', 'I', 'B'], legend=True, stacked=True)
    # Reverse words in sticks ( end to start because this is hebrew )
    reversed_sticks = []
    for stick in plt.xticks()[1]:
        reversed_sticks.append(stick.get_text()[::-1])
    plt.xticks(range(len(reversed_sticks)), reversed_sticks)
    plt.title('Top 10 words with highest frequency of O label')
    plt.ylabel('Frequency')
    plt.show()

    # Plot top 10 words with highest frequency of B label
    plt.figure(figsize=(10, 10))
    df.sort_values('B', ascending=False).head(10).plot.bar(y=['O', 'I', 'B'], legend=True, stacked=True)
    # Reverse words in sticks ( end to start because this is hebrew )
    reversed_sticks = []
    for stick in plt.xticks()[1]:
        reversed_sticks.append(stick.get_text()[::-1])
    plt.xticks(range(len(reversed_sticks)), reversed_sticks)
    plt.title('Top 10 words with highest frequency of B label')
    plt.ylabel('Frequency')
    plt.show()

    # Plot top 10 words with highest frequency of I label
    plt.figure(figsize=(10, 10))
    df.sort_values('I', ascending=False).head(10).plot.bar(y=['O', 'I', 'B'], legend=True,
                                                           stacked=True)
    # Reverse words in sticks ( end to start because this is hebrew )
    reversed_sticks = []
    for stick in plt.xticks()[1]:
        reversed_sticks.append(stick.get_text()[::-1])
    plt.xticks(range(len(reversed_sticks)), reversed_sticks)
    plt.title('Top 10 words with highest frequency of I label')
    plt.ylabel('Frequency')
    plt.show()


def get_sentence(data):
    full_sentence = ''
    for word in data:
        full_sentence += word + ' '
    return full_sentence


def get_metaphors(data, labels):
    metaphors = ''
    for word, label in zip(data, labels):
        if label == 0:
            continue
        else:
            metaphors += word + ' '
    return metaphors


def explore_data(show_plots):
    """
    Explore the data
    """
    create_word_cloud = False
    print_into_docs = False
    with open('../config/config.yaml') as f:
        training_args = Box(yaml.load(f, Loader=yaml.FullLoader))
    # Load data
    dataset = training_args.data_args.dataset
    full_dataset = load_dataset(dataset, split='all')
    train_dataset = load_dataset(dataset, split='train')
    test_dataset = load_dataset(dataset, split='test')
    validation_dataset = load_dataset(dataset, split='validation')

    seen_words = {}
    seen_metaphors = {}

    for example in train_dataset:
        words = example['data']
        labels = example['labels']
        for label, word in zip(labels, words):
            seen_words[word] = 1
            if label == 1 or label == 2:
                seen_metaphors[word] = 1

    print("Plotting statistics for a report")
    number_of_examples_train = len(train_dataset) + len(validation_dataset)
    number_of_examples_test = len(test_dataset)
    number_of_examples_total = number_of_examples_train + number_of_examples_test

    number_of_words_train = sum([len(example['data']) for example in train_dataset])
    number_of_words_train += sum([len(example['data']) for example in validation_dataset])
    number_of_words_test = sum([len(example['data']) for example in test_dataset])

    train_metaphor_number = sum([1 if label == 1 or label == 2 else 0 for example in train_dataset for label in example['labels']])
    train_metaphor_number += sum([1 if label == 1 or label == 2 else 0 for example in validation_dataset for label in example['labels']])
    train_metaphor_percentage = train_metaphor_number / number_of_words_train * 100

    test_metaphor_number = sum([1 if label == 1 or label == 2 else 0 for example in test_dataset for label in example['labels']])
    test_metaphor_percentage = test_metaphor_number / number_of_words_test * 100


    print("Number of examples in train: {}".format(number_of_examples_train))
    print("Number of examples in test: {}".format(number_of_examples_test))
    print("Number of examples in total: {}".format(number_of_examples_total))
    print("Percentage of metaphor in train: {}".format(train_metaphor_percentage))
    print("Percentage of metaphor in test: {}".format(test_metaphor_percentage))
    print("Number of words in train: {}".format(number_of_words_train))
    print("Number of words in test: {}".format(number_of_words_test))

    for dataset in [train_dataset, test_dataset, validation_dataset]:
        unseen_words_counter = 0
        seen_words_counter = 0
        unseen_metaphors_counter = 0
        seen_metaphors_counter = 0
        number_of_metaphors = 0
        number_of_non_metaphors = 0
        length_hist = []
        # Calculate statistics
        label_hist = np.zeros([3])
        word_dict = {}
        examples_with_metaphor = np.zeros([2])
        for example in dataset:
            bool_metaphor = False
            words = example['data']
            labels = example['labels']
            length_hist.append(len(words))
            for label, word in zip(labels, words):
                if word in seen_words:
                    seen_words_counter += 1
                else:
                    unseen_words_counter += 1
                if label == 0:
                    number_of_non_metaphors += 1
                else:
                    if word in seen_metaphors:
                        seen_metaphors_counter += 1
                    else:
                        unseen_metaphors_counter += 1
                    number_of_metaphors += 1
                    bool_metaphor = True
                if word not in word_dict:
                    word_dict[word] = np.zeros([3])
                    word_dict[word][label] = 1
                else:
                    word_dict[word][label] += 1
                label_hist[label] += 1
            if bool_metaphor:
                examples_with_metaphor[0] += 1
            else:
                examples_with_metaphor[1] += 1
        print("dataset: {}".format(dataset))
        print("Number of metaphors: {}".format(number_of_metaphors))
        print("Number of non metaphors: {}".format(number_of_non_metaphors))
        print("Number of metaphors out of all words: {}".format(
            number_of_metaphors / (number_of_metaphors + number_of_non_metaphors)))
        print("Number of unseen words: {}".format(unseen_words_counter))
        print("Number of seen words: {}".format(seen_words_counter))
        print("Number of unseen words out of all words: {}".format(
            unseen_words_counter / (seen_words_counter + unseen_words_counter)))
        print("Number of unseen metaphors: {}".format(unseen_metaphors_counter))
        print("Number of seen metaphors: {}".format(seen_metaphors_counter))
        print("Number of unseen metaphors out of all metaphors: {}".format(
            unseen_metaphors_counter / (seen_metaphors_counter + unseen_metaphors_counter)))
        print("-----------------------------------------------------")


        # plot length histogram
        plt.figure(figsize=(10, 5))
        plt.hist(length_hist, bins=100)
        plt.xlabel('Length')
        plt.ylabel('Count')
        plt.title('Length Distribution')
        if show_plots:
            plt.show()
        else:
            plt.close()



        # Plot statistics
        if create_word_cloud:
            comment_words = ''
            stopwords = set(STOPWORDS)

            # iterate data
            for example in dataset:
                words = example['data']
                labels = example['labels']
                for word, label in zip(words, labels):
                    if label > 0:
                        comment_words += " " + word + " "

            # reverse text ( end to start because this is hebrew )
            bidi_text = get_display(comment_words)

            wordcloud = WordCloud(width=800, height=800,
                                  background_color='white',
                                  stopwords=stopwords,
                                  font_path='../Font/FreeSans/FreeSansBold.ttf',
                                  min_font_size=10).generate(bidi_text)

            # plot the WordCloud image
            plt.figure(figsize=(8, 8), facecolor=None)
            plt.imshow(wordcloud)
            plt.axis("off")
            plt.tight_layout(pad=0)
            plt.savefig('wordcloud.png')
            if show_plots:
                plt.show()
            else:
                plt.close()

            plot_statistics(label_hist, word_dict, examples_with_metaphor)

        # Print the data with its labels
        example_df = pd.DataFrame(columns=['Full Sentence', 'Metaphor Words'])
        words_list, labels_list = [], []
        num_of_examples = 5
        np.random.seed(6324)
        indices = np.random.randint(0, len(dataset), num_of_examples)

        for index in indices:
            example = dataset[int(index)]
            example_df = example_df.append({'Full Sentence': get_sentence(example['data']),
                                            'Metaphor Words': get_metaphors(example['data'], example['labels'])},
                                           ignore_index=True)
            if print_into_docs:
                words_list.append(example['data'])
                labels_list.append(example['labels'])
        if print_into_docs:
            print_labels(words_list, labels_list, 'examples_with_labels_0922.docx')

        example_df = example_df[example_df.columns[::-1]]

        fig, ax = plt.subplots()
        ax.axis('tight')
        ax.axis('off')

        for example_i, (sentence, metaphors) in enumerate(zip(example_df['Full Sentence'], example_df['Metaphor Words'])):
            example_df['Full Sentence'][example_i] = get_display(sentence)
            example_df['Metaphor Words'][example_i] = get_display(metaphors)

        the_table = ax.table(cellText=example_df.values, colLabels=example_df.columns, loc='center')

        for i in range(0, len(example_df)):
            the_table.auto_set_column_width(i)

        the_table.auto_set_font_size(False)
        the_table.set_fontsize(8)
        the_table.scale(1, 1.5)

        # add a title
        ax.set_title('Examples from our dataset', fontsize=20, fontweight='bold')
        fig.tight_layout()

        plt.savefig('examples_with_labels.png')
        plt.show()

        fig, ax = plt.subplots()
        ax.axis('tight')
        ax.axis('off')
        # each row is a sentence aligned in the right
        # each metaphor is highlighted with yellow background

        # example_df = example_df.drop(['Metaphor Words'], axis=1)

        the_table = ax.table(cellText=example_df.values, colLabels=example_df.columns, loc='center')


        for i in range(0, len(example_df)):
            the_table.auto_set_column_width(i)

        the_table.auto_set_font_size(False)
        the_table.set_fontsize(18)
        the_table.scale(1, 2)

        # add a title
        ax.set_title('Examples from our dataset', fontsize=20, fontweight='bold')
        fig.tight_layout()

        plt.savefig('examples_with_labels.png')
        plt.show()


if __name__ == '__main__':
    show_plots = False
    explore_data(show_plots)
