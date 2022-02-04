from datasets import load_dataset, load_metric
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.enum.text import WD_COLOR_INDEX
from box import Box
import yaml
import os
os.environ['KMP_DUPLICATE_LIB_OK']='True'

label_names = {"O": 0, "B-metaphor": 1, "I-metaphor": 2}


def print_predictions(words_list, labels_list, predictions_list, filename):
    doc = docx.Document()
    doc.add_heading('Words with metaphors highlighted', 0)
    for index, (words, labels, predictions) in enumerate(zip(words_list, labels_list, predictions_list)):
        para = doc.add_paragraph('Example: ' + str(index) + '\n')
        for word, label, prediction in zip(words, labels, predictions):
            if label == 0 and prediction == 0:
                para.add_run(word + ' ')
            elif label != 0 and prediction != 0:  # green for correct prediction
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GREEN
            elif label != 0 and prediction == 0:  # red for metaphor that was not predicted
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.RED
            elif label == 0 and prediction != 0:  # pink for not metaphor that was predicted (wrongly)
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.PINK
        # save the document
        doc.save(filename)


def print_labels(words_list, labels_list, filename):
    doc = docx.Document()
    doc.add_heading('Words with metaphors highlighted', 0)
    for index, (words, labels) in enumerate(zip(words_list, labels_list)):
        para = doc.add_paragraph('Example: ' + str(index) + '\n')
        for word, label in zip(words, labels):
            if label == 0:
                para.add_run(word + ' ')
            elif label == 1:
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.BRIGHT_GREEN
            else:  # label == 2
                para.add_run(word + ' ').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GREEN
    # Now save the document
    doc.save(filename)


#  Plotting statistics
def plot_statisticts(label_hist, word_dict, examples_with_metaphor):
    plt.bar(range(3), label_hist)
    # show numbers on the bars
    for i in range(3):
        plt.text(i, label_hist[i], str(label_hist[i]))
    # change sticks names
    plt.xticks(range(3), [label for label in label_names])
    plt.xlabel('Label')
    plt.ylabel('Count')
    plt.title('Label Distribution')
    plt.show()

    # Number of examples with at least one metaphor
    plt.figure(figsize=(10, 5))
    plt.bar(range(2), examples_with_metaphor, align='center')
    # show numbers on the bars
    for i in range(2):
        plt.text(i, examples_with_metaphor[i], str(examples_with_metaphor[i]))
    plt.xticks(range(2), ['Without metaphor', 'With metaphor'])
    plt.xlabel('Label')
    plt.ylabel('Number of examples')
    plt.title('Number of examples with at least one label')
    plt.show()

    # create a dataframe from the dictionary
    df = pd.DataFrame.from_dict(word_dict, orient='index')
    df.rename(columns={0: "O", 1: "B", 2: "I"}, inplace=True)
    df.index.name = 'word'

    # Plot top 10 words with highest frequency of O label
    plt.figure(figsize=(10, 10))
    df.sort_values('O', ascending=False).head(10).plot.bar(y=['O', 'I', 'B'], legend=True, stacked=True)
    # Reverse words in sticks ( end to start because this is hebrew )
    reversed_sticks = []
    for stick in plt.xticks()[1]:
        reversed_sticks.append(stick.get_text()[::-1])
    plt.xticks(range(len(reversed_sticks)), reversed_sticks)
    plt.title('Top 10 words with highest frequency of O label')
    plt.ylabel('Frequency')
    plt.show()

    # Plot top 10 words with highest frequency of B label
    plt.figure(figsize=(10, 10))
    df.sort_values('B', ascending=False).head(10).plot.bar(y=['O', 'I', 'B'], legend=True, stacked=True)
    # Reverse words in sticks ( end to start because this is hebrew )
    reversed_sticks = []
    for stick in plt.xticks()[1]:
        reversed_sticks.append(stick.get_text()[::-1])
    plt.xticks(range(len(reversed_sticks)), reversed_sticks)
    plt.title('Top 10 words with highest frequency of B label')
    plt.ylabel('Frequency')
    plt.show()

    # Plot top 10 words with highest frequency of I label
    plt.figure(figsize=(10, 10))
    df.sort_values('I', ascending=False).head(10).plot.bar(y=['O', 'I', 'B'], legend=True,
                                                           stacked=True)
    # Reverse words in sticks ( end to start because this is hebrew )
    reversed_sticks = []
    for stick in plt.xticks()[1]:
        reversed_sticks.append(stick.get_text()[::-1])
    plt.xticks(range(len(reversed_sticks)), reversed_sticks)
    plt.title('Top 10 words with highest frequency of I label')
    plt.ylabel('Frequency')
    plt.show()


def explore_data():
    with open('config.yaml') as f:
        training_args = Box(yaml.load(f, Loader=yaml.FullLoader))
    # Load data
    dataset = load_dataset(training_args.data_args.dataset, split='all')

    # Calculate statistics
    label_hist = np.zeros([3])
    word_dict = {}
    examples_with_metaphor = np.zeros([2])
    for example in dataset:
        bool_metaphor = False
        words = example['data']
        labels = example['labels']
        for label, word in zip(labels, words):
            if word not in word_dict:
                word_dict[word] = np.zeros([3])
                word_dict[word][label] = 1
            else:
                word_dict[word][label] += 1
            label_hist[label] += 1
            if label > 0:
                bool_metaphor = True
        if bool_metaphor:
            examples_with_metaphor[0] += 1
        else:
            examples_with_metaphor[1] += 1

    #  Plotting statistics
    plot_statisticts(label_hist, word_dict, examples_with_metaphor)

    # Print the data with its labels
    words_list = []
    labels_list = []
    indices = np.arange(len(dataset))
    for index in indices:
        example = dataset[int(index)]
        words_list.append(example['data'])
        labels_list.append(example['labels'])
    print_labels(words_list, labels_list, 'examples_with_labels.docx')
